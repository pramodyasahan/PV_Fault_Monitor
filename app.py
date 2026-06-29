# ================================================================
#  PV PLANT FAULT MONITORING SYSTEM
#  Hybrid AI Architecture for Solar PV Fault Diagnosis
#  University of Moratuwa | Dept. of Electrical Engineering
# ================================================================

import os
import io
import json
import re
import pickle
import warnings
from datetime import datetime

import numpy as np
import pandas as pd
import plotly.graph_objects as go
import streamlit as st
import torch
import torch.nn as nn

warnings.filterwarnings("ignore")

# ── PAGE CONFIG ──────────────────────────────────────────────────
st.set_page_config(
    page_title="PV Fault Monitor",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── CONSTANTS ────────────────────────────────────────────────────
BASE     = os.path.dirname(os.path.abspath(__file__))
MODELS   = os.path.join(BASE, "models")
SAMPLES  = os.path.join(BASE, "sample_data")
N_SEQ    = 50
WIN_Z2   = 32

Z1_S1_LABELS = {
    0: "Healthy",
    1: "Degradation (DG)",
    2: "Short-Circuit (SC)",
    3: "Open-Circuit (OC)",
}
Z1_S2_LABELS = {
    0: "SC — Array 1  |  R_LL = 0.05 Ω (Severe)",
    1: "SC — Array 1  |  R_LL = 10 Ω",
    2: "SC — Array 1  |  R_LL = 20 Ω",
    3: "SC — Array 1  |  R_LL = 30 Ω (Mild)",
    4: "SC — Array 2  |  R_LL = 0.05 Ω (Severe)",
    5: "SC — Array 2  |  R_LL = 10 Ω",
    6: "SC — Array 2  |  R_LL = 20 Ω",
    7: "SC — Array 2  |  R_LL = 30 Ω (Mild)",
    8: "Open-Circuit — Array 1",
    9: "Open-Circuit — Array 2",
}
Z2_LABELS = {
    0: "Healthy",
    1: "T1 Open-Circuit Fault",
    2: "T2 Open-Circuit Fault",
    3: "T3 Open-Circuit Fault",
    4: "T4 Open-Circuit Fault",
    5: "T5 Open-Circuit Fault",
    6: "T6 Open-Circuit Fault",
}

# ── CUSTOM CSS ───────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;600;700;800&display=swap');

:root {
    --term-bg: #05080d;
    --term-panel: #0a1018;
    --term-panel-2: #0e1624;
    --term-line: #17314a;
    --term-green: #22c55e;
    --term-cyan: #22d3ee;
    --term-amber: #f59e0b;
    --term-red: #ef4444;
    --term-text: #dbeafe;
    --term-muted: #7dd3fc;
}

[data-testid="stAppViewContainer"] {
    background:
        radial-gradient(circle at top left, rgba(34,211,238,0.10), transparent 28%),
        radial-gradient(circle at top right, rgba(34,197,94,0.08), transparent 24%),
        linear-gradient(180deg, #030712 0%, #05080d 100%);
}
[data-testid="stHeader"]  { background: transparent; }
[data-testid="stSidebar"] { background: #070b12; }
.block-container { padding-top: 1.1rem; }
html, body, [class*="css"] { font-family: 'JetBrains Mono', monospace; }

.main-title {
    text-align: center;
    font-family: 'JetBrains Mono', monospace;
    font-size: 29px;
    font-weight: 800;
    color: #d1fae5;
    letter-spacing: 0.6px;
    text-shadow: 0 0 16px rgba(34,197,94,0.36);
    margin-bottom: 4px;
}
.main-sub {
    text-align: center;
    font-family: 'JetBrains Mono', monospace;
    font-size: 11px;
    color: #67e8f9;
    margin-bottom: 16px;
}
.zone-title {
    font-family: 'JetBrains Mono', monospace;
    font-size: 13px;
    font-weight: 800;
    color: #cffafe;
    background: linear-gradient(90deg, rgba(34,211,238,0.16), rgba(34,197,94,0.06));
    border: 1px solid #164e63;
    border-left: 4px solid #22d3ee;
    border-radius: 8px;
    padding: 9px 10px;
    margin-bottom: 10px;
    text-transform: uppercase;
    letter-spacing: 0.5px;
}
.zone-title::after {
    content: " ONLINE";
    float: right;
    color: #22c55e;
    font-size: 10px;
    font-weight: 800;
    letter-spacing: 1px;
}
.demo-caption {
    color: #67e8f9;
    font-size: 10px;
    font-weight: 800;
    letter-spacing: 0.7px;
    text-transform: uppercase;
    margin: 4px 0 8px 0;
}
.badge {
    border-radius: 8px;
    padding: 8px 12px;
    font-family: 'JetBrains Mono', monospace;
    font-size: 12px;
    font-weight: 800;
    text-align: center;
    margin: 6px 0;
    display: block;
    box-shadow: inset 0 0 18px rgba(255,255,255,0.03), 0 0 18px rgba(34,211,238,0.05);
}
.metric-row {
    background: #08111d;
    border-radius: 7px;
    padding: 8px 11px;
    margin: 5px 0;
    font-family: 'JetBrains Mono', monospace;
    font-size: 11px;
    border-left: 3px solid #22d3ee;
    color: #cbd5e1;
}
.prob-wrap { margin: 5px 0; }
.prob-label {
    display: flex;
    justify-content: space-between;
    font-family: 'JetBrains Mono', monospace;
    font-size: 10px;
    color: #93c5fd;
    margin-bottom: 3px;
}
.prob-track {
    height: 5px;
    background: #07111d;
    border: 1px solid #12263d;
    border-radius: 3px;
    overflow: hidden;
}
.terminal-card {
    background: linear-gradient(180deg, rgba(8,17,29,0.98), rgba(3,7,18,0.98));
    border: 1px solid #164e63;
    border-radius: 12px;
    padding: 12px 13px;
    margin: 8px 0;
    box-shadow: 0 0 22px rgba(34,211,238,0.07), inset 0 0 22px rgba(15,23,42,0.35);
}
.terminal-card-title {
    color: #22d3ee;
    font-family: 'JetBrains Mono', monospace;
    font-size: 11px;
    font-weight: 800;
    text-transform: uppercase;
    letter-spacing: 0.8px;
    margin-bottom: 8px;
}
.terminal-line {
    display: flex;
    justify-content: space-between;
    gap: 12px;
    padding: 4px 0;
    border-bottom: 1px dashed rgba(34,211,238,0.12);
    font-family: 'JetBrains Mono', monospace;
    font-size: 11px;
}
.terminal-line:last-child { border-bottom: none; }
.terminal-key { color: #7dd3fc; }
.terminal-value { color: #d1fae5; font-weight: 800; text-align: right; }
div[data-testid="stMetric"] {
    background: #07111d;
    border: 1px solid #12324a;
    border-radius: 10px;
    padding: 7px 9px;
}
div[data-testid="stButton"] > button,
div[data-testid="stDownloadButton"] > button {
    border-radius: 8px;
    font-family: 'JetBrains Mono', monospace;
    font-weight: 800;
    border: 1px solid #164e63;
    background: linear-gradient(180deg, #0e7490, #064e3b);
    color: #ecfeff;
}
hr { border-color: #12324a !important; }

/* ── Analysis Section (modern, clean) ── */
.analysis-wrapper {
    background: linear-gradient(160deg, rgba(8,17,29,0.97), rgba(3,9,20,0.98));
    border: 1px solid #1e3a5f;
    border-radius: 16px;
    padding: 24px 26px;
    margin-top: 4px;
    box-shadow: 0 0 32px rgba(34,211,238,0.05);
}
.analysis-top-row {
    display: flex;
    align-items: center;
    justify-content: space-between;
    margin-bottom: 18px;
    padding-bottom: 12px;
    border-bottom: 1px solid #1e3a5f;
}
.analysis-title {
    font-family: 'JetBrains Mono', monospace;
    font-size: 15px;
    font-weight: 800;
    color: #e0f2fe;
    letter-spacing: 0.4px;
}
.analysis-subtitle {
    font-family: 'JetBrains Mono', monospace;
    font-size: 10px;
    color: #475569;
    margin-top: 3px;
}
.report-card-outer {
    background: #020a14;
    border: 1px solid #0f2744;
    border-radius: 12px;
    padding: 22px 24px;
    margin-top: 14px;
    max-height: 540px;
    overflow-y: auto;
}
.report-section-heading {
    font-family: 'JetBrains Mono', monospace;
    font-size: 11px;
    font-weight: 800;
    color: #22d3ee;
    text-transform: uppercase;
    letter-spacing: 1px;
    margin: 16px 0 8px 0;
    padding-bottom: 5px;
    border-bottom: 1px solid #0f2744;
}
.report-section-heading:first-child { margin-top: 0; }
.report-body-text {
    font-family: 'JetBrains Mono', monospace;
    font-size: 11.5px;
    color: #cbd5e1;
    line-height: 1.85;
    margin: 0 0 6px 0;
}
.report-bullet {
    font-family: 'JetBrains Mono', monospace;
    font-size: 11.5px;
    color: #94a3b8;
    line-height: 1.85;
    padding-left: 14px;
    position: relative;
    margin: 3px 0;
}
.report-bullet::before {
    content: "›";
    position: absolute;
    left: 0;
    color: #22d3ee;
    font-weight: 800;
}
.report-bold { color: #dbeafe; font-weight: 800; }
.export-panel {
    background: rgba(8,17,29,0.85);
    border: 1px solid #1e3a5f;
    border-radius: 14px;
    padding: 20px;
    height: 100%;
}
.export-panel-title {
    font-family: 'JetBrains Mono', monospace;
    font-size: 11px;
    font-weight: 800;
    color: #67e8f9;
    text-transform: uppercase;
    letter-spacing: 1px;
    margin-bottom: 14px;
    padding-bottom: 8px;
    border-bottom: 1px solid #1e3a5f;
}
.status-chip {
    display: inline-block;
    padding: 3px 9px;
    border-radius: 999px;
    font-family: 'JetBrains Mono', monospace;
    font-size: 10px;
    font-weight: 800;
    margin: 2px 3px 2px 0;
}
</style>
""", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════════
# MODEL ARCHITECTURES
# ════════════════════════════════════════════════════════════════

class LSTMClassifier(nn.Module):
    def __init__(self, input_size, hidden, n_layers, n_classes, dropout=0.2):
        super().__init__()
        self.lstm = nn.LSTM(input_size, hidden, n_layers, batch_first=True,
                            dropout=dropout if n_layers > 1 else 0.0)
        self.head = nn.Sequential(
            nn.Dropout(dropout),
            nn.Linear(hidden, 64), nn.ReLU(),
            nn.Linear(64, n_classes),
        )
    def forward(self, x):
        out, _ = self.lstm(x)
        return self.head(out[:, -1, :])


class Zone2CNN1D(nn.Module):
    def __init__(self):
        super().__init__()
        self.conv_stack = nn.Sequential(
            nn.Conv1d(4,  32, 3, padding=1), nn.BatchNorm1d(32),  nn.ReLU(), nn.MaxPool1d(2),
            nn.Conv1d(32, 64, 3, padding=1), nn.BatchNorm1d(64),  nn.ReLU(), nn.MaxPool1d(2),
            nn.Conv1d(64,128, 3, padding=1), nn.BatchNorm1d(128), nn.ReLU(), nn.AdaptiveAvgPool1d(1),
        )
        self.head = nn.Linear(128, 7)
    def forward(self, x):
        return self.head(self.conv_stack(x).squeeze(-1))


# ════════════════════════════════════════════════════════════════
# MODEL LOADING
# ════════════════════════════════════════════════════════════════

@st.cache_resource(show_spinner="Loading AI models…")
def load_models() -> dict:
    m = {}
    try:
        s1 = LSTMClassifier(5, 128, 2, 4)
        s1.load_state_dict(torch.load(os.path.join(MODELS,"zone1_stage1_lstm.pth"), map_location="cpu"))
        s1.eval(); m["z1s1"] = s1
        with open(os.path.join(MODELS,"zone1_stage1_scaler.pkl"),"rb") as f: m["z1s1_sc"] = pickle.load(f)

        s2 = LSTMClassifier(8, 128, 2, 10)
        s2.load_state_dict(torch.load(os.path.join(MODELS,"zone1_stage2_lstm.pth"), map_location="cpu"))
        s2.eval(); m["z1s2"] = s2
        with open(os.path.join(MODELS,"zone1_stage2_scaler.pkl"),"rb") as f: m["z1s2_sc"] = pickle.load(f)

        z2 = Zone2CNN1D()
        z2.load_state_dict(torch.load(os.path.join(MODELS,"zone2_model.pt"), map_location="cpu"))
        z2.eval(); m["z2"] = z2
        with open(os.path.join(MODELS,"zone2_scaler.pkl"),"rb") as f: m["z2_sc"] = pickle.load(f)

        with open(os.path.join(MODELS,"zone3_model.pkl"),"rb") as f: m["z3"] = pickle.load(f)
        with open(os.path.join(MODELS,"zone3_scaler.pkl"),"rb") as f: m["z3_sc"] = pickle.load(f)

        m["ok"] = True
    except Exception as e:
        m["ok"] = False; m["err"] = str(e)
    return m


# ════════════════════════════════════════════════════════════════
# PREPROCESSING
# ════════════════════════════════════════════════════════════════

def prep_z1_sys(df, scaler):
    df = df.sort_values("V").reset_index(drop=True)
    idx = np.linspace(0, len(df) - 1, N_SEQ).astype(int)
    seq = df.iloc[idx][["V", "I", "P", "Irr", "T"]].values.astype(np.float32)
    return scaler.transform(seq.reshape(-1, 5)).reshape(1, N_SEQ, 5)

def prep_z1_arr(df, scaler):
    df = df.sort_values("V_1").reset_index(drop=True)
    idx = np.linspace(0, len(df) - 1, N_SEQ).astype(int)
    seq = df.iloc[idx][["V_1","I_1","P_1","V_2","I_2","P_2","Irr","T"]].values.astype(np.float32)
    return scaler.transform(seq.reshape(-1, 8)).reshape(1, N_SEQ, 8)

def prep_z2_best_window(df, scaler, model):
    cols = ["I_alpha", "I_beta", "Irradiance", "Temperature"]
    data = df[cols].values.astype(np.float32)
    n    = len(data)
    best_probs, best_x = None, None
    for start in range(max(1, n - WIN_Z2 + 1)):
        win    = data[start:start + WIN_Z2]
        x      = torch.tensor(scaler.transform(win).T[np.newaxis], dtype=torch.float32)
        with torch.no_grad():
            probs = torch.softmax(model(x), 1)[0].numpy()
        if best_probs is None or float(probs.max()) > float(best_probs.max()):
            best_probs = probs; best_x = x
    return best_x, best_probs

def prep_z3(df, scaler):
    cols = ["Irradiance", "T_ambient", "I_rms_high", "V_avg", "V_ripple"]
    return scaler.transform(df[cols].values[:1].astype(np.float32))


# ════════════════════════════════════════════════════════════════
# INFERENCE
# ════════════════════════════════════════════════════════════════

@torch.no_grad()
def run_z1(df_sys, df_arr, m):
    x     = torch.tensor(prep_z1_sys(df_sys, m["z1s1_sc"]))
    logit = m["z1s1"](x)
    cls   = int(logit.argmax(1))
    prob  = torch.softmax(logit, 1)[0].numpy()
    s2_cls, s2_lbl, s2_info = None, None, {}
    if cls in [2, 3] and df_arr is not None:
        x2     = torch.tensor(prep_z1_arr(df_arr, m["z1s2_sc"]))
        logit2 = m["z1s2"](x2)
        s2_cls = int(logit2.argmax(1))
        s2_lbl = Z1_S2_LABELS[s2_cls]
        s2_info = parse_z1_stage2(s2_lbl, s2_cls)
    return {"cls": cls, "label": Z1_S1_LABELS[cls],
            "probs": prob, "s2_cls": s2_cls, "s2_label": s2_lbl, **s2_info}

def run_z2(df, m):
    _, prob = prep_z2_best_window(df, m["z2_sc"], m["z2"])
    cls     = int(prob.argmax())
    return {"cls": cls, "label": Z2_LABELS[cls], "probs": prob}

def run_z3(df, m):
    X_s   = prep_z3(df, m["z3_sc"])
    esr   = float(np.clip(m["z3"].predict(X_s)[0], 0.15, 0.40))
    deg   = esr_to_degradation_percentage(esr)
    fault = esr >= 0.30
    row   = df.iloc[0]
    return {
        "esr": esr, "deg": deg, "fault": fault,
        "label": zone3_health_label(esr, deg),
        "irr":  float(row.get("Irradiance", 0)),
        "tamb": float(row.get("T_ambient", 0)),
        "irms": float(row.get("I_rms_high", 0)),
        "vavg": float(row.get("V_avg", 0)),
        "vrip": float(row.get("V_ripple", 0)),
    }


# ════════════════════════════════════════════════════════════════
# VISUALISATIONS
# ════════════════════════════════════════════════════════════════

_PLOT_CFG = {"displayModeBar": False}
_BG_PANEL = "#131720"
_AXIS_COL = "#475569"
_GRID_COL = "#1a2133"

def _base_layout(**kwargs):
    return dict(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor=_BG_PANEL,
                height=195, margin=dict(l=38, r=38, t=10, b=36),
                font=dict(color=_AXIS_COL, size=10), **kwargs)

def plot_iv(df):
    fig = go.Figure()
    if df is None or not all(c in df.columns for c in ["V","I","P"]):
        fig.add_annotation(text="Invalid or missing columns.<br>Expected: V, I, P",
                           xref="paper", yref="paper", x=0.5, y=0.5, showarrow=False,
                           font=dict(size=12, color="#ef4444"))
        fig.update_layout(**_base_layout()); return fig
    fig.add_trace(go.Scatter(x=df["V"], y=df["I"], mode="lines", name="I–V",
                             line=dict(color="#3b82f6", width=2.5)))
    fig.add_trace(go.Scatter(x=df["V"], y=df["P"], mode="lines", name="P–V",
                             line=dict(color="#f59e0b", width=2, dash="dash"), yaxis="y2"))
    fig.update_layout(**_base_layout(
        xaxis=dict(title="Voltage (V)", gridcolor=_GRID_COL, color=_AXIS_COL),
        yaxis=dict(title="Current (A)", gridcolor=_GRID_COL, color=_AXIS_COL),
        yaxis2=dict(title="Power (W)", overlaying="y", side="right", color="#f59e0b"),
        legend=dict(x=0.02, y=0.97, bgcolor="rgba(0,0,0,0)", font=dict(size=9)),
    ))
    return fig

def plot_ab(df):
    fig = go.Figure()
    if df is None or not all(c in df.columns for c in ["I_alpha","I_beta"]):
        fig.add_annotation(text="Invalid or missing columns.<br>Expected: I_alpha, I_beta",
                           xref="paper", yref="paper", x=0.5, y=0.5, showarrow=False,
                           font=dict(size=12, color="#ef4444"))
        fig.update_layout(**_base_layout()); return fig
    fig.add_trace(go.Scatter(x=df["I_alpha"], y=df["I_beta"], mode="lines",
                             line=dict(color="#a855f7", width=2.5)))
    fig.update_layout(**_base_layout(
        xaxis=dict(title="Iα (A)", gridcolor=_GRID_COL, color=_AXIS_COL, scaleanchor="y"),
        yaxis=dict(title="Iβ (A)", gridcolor=_GRID_COL, color=_AXIS_COL),
    ))
    return fig

def plot_gauge(esr, deg):
    c = "#ef4444" if esr >= 0.30 else ("#f59e0b" if deg >= 70 else "#22c55e")
    fig = go.Figure(go.Indicator(
        mode="gauge+number", value=round(deg, 1),
        title={"text": f"ESR = {esr:.4f} Ω", "font": {"size": 11, "color": _AXIS_COL}},
        number={"suffix": "%", "font": {"size": 30, "color": c}},
        gauge={
            "axis": {"range": [0, 100], "tickcolor": _AXIS_COL},
            "bar":  {"color": c}, "bgcolor": _BG_PANEL, "bordercolor": _GRID_COL,
            "steps": [
                {"range": [0,  50], "color": "rgba(34,197,94,0.10)"},
                {"range": [50, 80], "color": "rgba(245,158,11,0.10)"},
                {"range": [80,100], "color": "rgba(239,68,68,0.14)"},
            ],
            "threshold": {"line": {"color": "#ef4444", "width": 3}, "thickness": 0.8, "value": 100},
        },
    ))
    fig.update_layout(paper_bgcolor="rgba(0,0,0,0)", height=195,
                      margin=dict(l=16,r=16,t=28,b=8), font=dict(color="#e2e8f0"))
    return fig


# ════════════════════════════════════════════════════════════════
# UI HELPERS
# ════════════════════════════════════════════════════════════════

def badge(text, color):
    return (f'<div class="badge" style="background:{color}22;border:1px solid {color};'
            f'color:{color}">{text}</div>')

def prob_bars(labels, probs, pred_cls):
    html = ""
    for i, name in labels.items():
        p, hi = float(probs[i]), i == pred_cls
        col = "#3b82f6" if hi else "#334155"
        fw  = "700" if hi else "400"
        w   = int(p * 100)
        html += (f'<div class="prob-wrap"><div class="prob-label"><span>{name}</span>'
                 f'<span style="color:{col};font-weight:{fw}">{p*100:.1f}%</span></div>'
                 f'<div class="prob-track"><div style="width:{w}%;height:100%;'
                 f'background:{col};border-radius:2px"></div></div></div>')
    return html

def parse_z1_stage2(label, cls):
    if label is None: return {}
    info = {"stage2_raw": label, "stage2_cls": cls, "stage2_fault_type": None,
            "stage2_array": None, "r_ll": None, "r_ll_float": None, "sc_severity": None}
    if "Array 1" in label:  info["stage2_array"] = "Array 1"
    elif "Array 2" in label: info["stage2_array"] = "Array 2"
    if label.startswith("SC"):
        info["stage2_fault_type"] = "Short-Circuit"
        try:
            raw = label.split("R_LL =", 1)[1].split("Ω", 1)[0].strip()
            info["r_ll"] = f"{raw} Ω"; info["r_ll_float"] = float(raw)
        except Exception:
            info["r_ll"] = "Unavailable"
        r = info["r_ll_float"]
        if r is not None:
            if r <= 0.1: info["sc_severity"] = "Critical / near-solid short"
            elif r <= 10: info["sc_severity"] = "High severity short"
            elif r <= 20: info["sc_severity"] = "Moderate severity short"
            else: info["sc_severity"] = "Mild short-circuit path"
    elif label.startswith("Open-Circuit"):
        info["stage2_fault_type"] = "Open-Circuit"
        info["sc_severity"] = "Not applicable"
    return info

def terminal_panel(title, rows, color="#22d3ee"):
    html = f'<div class="terminal-card" style="border-color:{color}"><div class="terminal-card-title">▣ {title}</div>'
    for k, v in rows.items():
        html += (f'<div class="terminal-line"><span class="terminal-key">{k}</span>'
                 f'<span class="terminal-value">{v}</span></div>')
    return html + '</div>'

def esr_to_degradation_percentage(esr):
    return float(np.clip((esr - 0.15) / (0.30 - 0.15) * 100, 0, 100))

def zone3_health_label(esr, deg):
    if esr >= 0.30: return "Fault — ESR threshold exceeded / 100% degraded"
    if esr > 0.15:  return f"Ageing — {deg:.1f}% degraded"
    return "Healthy — 0.0% degraded"

def zone_status(r, zone):
    if r is None: return "gray", "⚪ Awaiting Data"
    if zone == 1:
        m = {0:("#22c55e","✅ Healthy"),1:("#f59e0b","⚠️ Degradation"),
             2:("#ef4444","🔴 Short-Circuit"),3:("#ef4444","🔴 Open-Circuit")}
        return m[r["cls"]]
    if zone == 2:
        return ("#22c55e","✅ Healthy") if r["cls"]==0 else ("#ef4444","🔴 IGBT Fault")
    if r["fault"]:       return "#ef4444","🔴 Capacitor Degraded"
    if r["deg"] >= 70:   return "#f59e0b","⚠️ Capacitor Ageing"
    return "#22c55e","✅ Capacitor Healthy"

def clean_df(df):
    if df is None: return None
    df.columns = [c.strip().replace('\ufeff','') for c in df.columns]
    mapping = {
        "v":"V","i":"I","p":"P","irr":"Irr","t":"T",
        "v_1":"V_1","i_1":"I_1","p_1":"P_1","v_2":"V_2","i_2":"I_2","p_2":"P_2",
        "i_alpha":"I_alpha","i_beta":"I_beta",
        "irradiance":"Irradiance","temperature":"Temperature",
        "t_ambient":"T_ambient","i_rms_high":"I_rms_high",
        "v_avg":"V_avg","v_ripple":"V_ripple"
    }
    rn = {c: mapping[c.lower()] for c in df.columns if c.lower() in mapping and c != mapping[c.lower()]}
    return df.rename(columns=rn) if rn else df

def sample(fname):
    p = os.path.join(SAMPLES, fname)
    return clean_df(pd.read_csv(p, encoding="utf-8-sig")) if os.path.exists(p) else None

def _confidence(result):
    if not result or "probs" not in result: return "N/A"
    try: return f"{float(np.max(result['probs']))*100:.1f}%"
    except: return "N/A"


# ════════════════════════════════════════════════════════════════
# LLM + LOCAL FALLBACK
# ════════════════════════════════════════════════════════════════

def _sanitize_report_text(text: str) -> str:
    text = "" if text is None else str(text)
    # Remove square-bracket urgency labels only
    text = re.sub(
        r"\[(?:URGENT|CRITICAL|WITHIN\s*48H|SCHEDULED|IMMEDIATE|PRIORITY\s*\d+)\]\s*",
        "", text, flags=re.IGNORECASE
    )
    return text.strip()

def _local_report(r1, r2, r3):
    sections = []
    faults = []
    if r1 and r1.get("cls") != 0: faults.append(f"Zone 1 PV Array ({r1['label']})")
    if r2 and r2.get("cls") != 0: faults.append(f"Zone 2 IGBT Inverter ({r2['label']})")
    if r3 and r3.get("fault"):    faults.append(f"Zone 3 DC-Link Capacitor ({r3['label']})")
    sections.append("**1. Executive System Health Summary**")
    if faults:
        sections.append(f"The hybrid diagnostic system detected {len(faults)} abnormal condition(s): "
                        f"{', '.join(faults)}. Results should be verified using site measurements and standard electrical safety procedures.")
    else:
        sections.append("All evaluated zones report nominal operating conditions. No PV array, IGBT inverter, or DC-link capacitor fault was detected for the supplied data.")
    sections.append("\n**2. Evidence From AI Models**")
    if r1:
        line = f"- **Zone 1 — PV Array LSTM:** predicted {r1['label']} with confidence {_confidence(r1)}."
        if r1.get("s2_label"): line += f" Stage 2 output: {r1['s2_label']}."
        sections.append(line)
    else: sections.append("- **Zone 1 — PV Array LSTM:** no input data supplied.")
    if r2:
        sections.append(f"- **Zone 2 — CNN-1D IGBT model:** predicted {r2['label']} with confidence {_confidence(r2)}.")
    else: sections.append("- **Zone 2 — CNN-1D IGBT model:** no input data supplied.")
    if r3:
        sections.append(f"- **Zone 3 — XGBoost ESR estimator:** predicted ESR = {r3['esr']:.4f} Ω, degradation = {r3['deg']:.1f}%.")
    else: sections.append("- **Zone 3 — XGBoost ESR estimator:** no input data supplied.")
    sections.append("\n**3. Fault Physics Interpretation**")
    if r1 and r1.get("cls") == 2:
        sections.append("- **PV short-circuit fault:** line-to-line short-circuit detected. Creates abnormal current paths and increases thermal stress.")
        if r1.get("s2_label"):
            sections.append(f"  Localised to {r1.get('stage2_array','Unknown')}. Estimated R_LL = {r1.get('r_ll','Unavailable')}. Severity: {r1.get('sc_severity','Unavailable')}. Lower R_LL indicates higher electrical stress.")
    elif r1 and r1.get("cls") == 3:
        sections.append("- **PV open-circuit fault:** string current path interrupted. R_LL is not applicable.")
        if r1.get("s2_label"): sections.append(f"  Localised to {r1.get('stage2_array','Unknown')}.")
    elif r1 and r1.get("cls") == 1:
        sections.append("- **PV array degradation:** reduced I-V performance consistent with cell ageing, mismatch, or soiling.")
    else: sections.append("- Zone 1: No PV array fault detected.")
    if r2 and r2.get("cls") != 0:
        sections.append(f"- **IGBT inverter fault:** {r2.get('label','Unknown')} distorts the alpha-beta current trajectory and increases harmonic content.")
    elif r2: sections.append("- Zone 2: Alpha-beta trajectory consistent with healthy inverter operation.")
    if r3 and r3.get("fault"):
        sections.append(f"- **DC-link capacitor fault:** ESR = {r3['esr']:.4f} Ω exceeds the 0.30 Ω fault threshold. Higher ESR increases ripple voltage and heat generation.")
    elif r3 and r3.get("esr",0) > 0.15:
        sections.append(f"- **DC-link capacitor ageing:** ESR in the 0.15–0.30 Ω ageing band. Degradation index: {r3['deg']:.1f}%.")
    else: sections.append("- Zone 3: ESR at healthy baseline.")
    sections.append("\n**4. Engineering Response Guidance**")
    guidance = []
    if r1 and r1.get("cls") == 2:
        guidance.append(f"Inspect {r1.get('stage2_array','the affected array')} for insulation damage, connector faults, and bypass-diode issues. Verify R_LL = {r1.get('r_ll','Unavailable')} using electrical tests.")
    elif r1 and r1.get("cls") == 3:
        guidance.append(f"Inspect {r1.get('stage2_array','the affected array')} for string discontinuity, loose connectors, and blown fuses.")
    elif r1 and r1.get("cls") == 1:
        guidance.append("Inspect PV modules for soiling, shading, and thermal anomalies. Compare with baseline I-V curves.")
    if r2 and r2.get("cls") != 0:
        guidance.append(f"Verify predicted {r2.get('label')} by checking gate-drive signals, switch continuity, and inverter thermal condition.")
    if r3 and r3.get("fault"):
        guidance.append("Verify capacitor ESR and ripple voltage using instruments, then evaluate replacement according to site maintenance procedure.")
    elif r3 and r3.get("deg",0) >= 70:
        guidance.append("Track capacitor ESR trend and plan component-level inspection.")
    if not guidance: guidance.append("Continue routine condition monitoring and compare against the healthy operating baseline.")
    for item in guidance: sections.append(f"- {item}")
    sections.append("\n**5. Hybrid Architecture Note**")
    sections.append("LSTM captures sequential I-V sweep behaviour for PV arrays. CNN-1D captures geometric patterns in alpha-beta inverter current trajectories. XGBoost handles tabular ESR-related capacitor features.")
    sections.append("\n*Report generated by local engineering analysis engine.*")
    return _sanitize_report_text("\n".join(sections))

def gpt_explain(r1, r2, r3):
    try:
        api_key = st.secrets.get("OPENAI_API_KEY","")
        if not api_key: return _local_report(r1, r2, r3)
        from openai import OpenAI
        client = OpenAI(api_key=api_key, timeout=8.0)
        findings = []
        if r1:
            line = f"Zone 1 (PV Array): **{r1['label']}** | confidence={_confidence(r1)}"
            if r1.get("s2_label"): line += f" -> Stage 2: {r1['s2_label']}"
            if r1.get("stage2_fault_type") == "Short-Circuit":
                line += f" | faulty_array={r1.get('stage2_array')}, R_LL={r1.get('r_ll')}, severity={r1.get('sc_severity')}"
            elif r1.get("stage2_fault_type") == "Open-Circuit":
                line += f" | faulty_array={r1.get('stage2_array')}; R_LL=not applicable"
            findings.append(line)
        if r2: findings.append(f"Zone 2 (IGBT Inverter): **{r2['label']}** | confidence={_confidence(r2)}")
        if r3: findings.append(f"Zone 3 (DC-Link Capacitor): **{r3['label']}** | ESR={r3['esr']:.4f} Ω | Degradation={r3['deg']:.1f}%")
        findings_str = "\n".join(f"- {f}" for f in findings) if findings else "- No faults detected."
        prompt = f"""You are a senior solar PV plant maintenance engineer writing the technical body of an engineering diagnostic report.

The report already contains an executive summary table and zone output table. Write ONLY the technical analysis using these four bold section titles — do NOT use markdown ## headers, do NOT repeat an executive summary:

**Fault Physics Interpretation**
Write what each detected fault means physically and its impact on system performance.

**Engineering Response Guidance**
Write specific inspection and verification steps for each fault using neutral language: inspect, isolate, verify, compare with baseline, plan component-level assessment.

**Confidence Assessment**
Comment on model confidence levels and whether results warrant field verification.

**Hybrid Architecture Note**
One short paragraph explaining why different model types are used for each zone (LSTM for IV curves, CNN-1D for alpha-beta trajectories, XGBoost for tabular ESR features).

DIAGNOSED FAULTS:
{findings_str}

Rules:
- Bold section titles only, no ## markdown headers, no numbered sections
- No urgency bracket labels
- When Zone 1 is SC: state the faulty array, R_LL value, severity — lower R_LL means lower impedance and higher electrical stress
- When Zone 1 is OC: state the faulty array — R_LL is not applicable for open-circuit
- When Zone 3 ESR is between 0.15 and 0.30 Ω: describe as a health/ageing index result, not a hard fault
- When Zone 3 ESR is >= 0.30 Ω: describe as fault-level degradation
- Do not invent numerical values not supplied above
- Under 380 words to fit everything into 1 page"""
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role":"user","content":prompt}],
            max_tokens=650, temperature=0.2,
        )
        return _sanitize_report_text(resp.choices[0].message.content)
    except Exception:
        return _local_report(r1, r2, r3)


# ════════════════════════════════════════════════════════════════
# REPORT RENDERING — clean HTML for UI display
# ════════════════════════════════════════════════════════════════

def render_report_html(text: str) -> str:
    """Convert markdown-flavoured report text to clean styled HTML for display."""
    if not text: return ""
    html_parts = []
    for line in text.split("\n"):
        s = line.strip()
        if not s:
            html_parts.append("<div style='height:6px'></div>")
            continue
        # Section headings: **1. Title** or **Title**
        m = re.match(r"^\*\*(\d+\.\s*.+?)\*\*$", s)
        if m:
            html_parts.append(f'<div class="report-section-heading">{m.group(1)}</div>')
            continue
        # Italic footer
        if s.startswith("*") and s.endswith("*") and not s.startswith("**"):
            inner = s.strip("*")
            html_parts.append(f'<div style="font-size:10px;color:#475569;margin-top:10px;font-style:italic">{inner}</div>')
            continue
        # Bullet points
        if s.startswith("- "):
            content = re.sub(r"\*\*(.*?)\*\*", r'<span class="report-bold">\1</span>', s[2:])
            html_parts.append(f'<div class="report-bullet">{content}</div>')
            continue
        # Sub-bullet (starts with spaces then -)
        if re.match(r"^\s{2,}-\s", line):
            content = re.sub(r"\*\*(.*?)\*\*", r'<span class="report-bold">\1</span>', s.lstrip("- "))
            html_parts.append(f'<div class="report-bullet" style="margin-left:18px;font-size:10.5px">{content}</div>')
            continue
        # Regular paragraph
        content = re.sub(r"\*\*(.*?)\*\*", r'<span class="report-bold">\1</span>', s)
        html_parts.append(f'<div class="report-body-text">{content}</div>')
    return "".join(html_parts)


# ════════════════════════════════════════════════════════════════
# EXPORT — Markdown
# ════════════════════════════════════════════════════════════════

def make_markdown_report(r1, r2, r3, llm_text: str) -> str:
    ts  = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    sep = "---"
    lines = [
        "# PV Plant Fault Diagnosis Report",
        "",
        f"**Generated:** {ts}  ",
        "**System:** 5 kW Grid-Connected PV Installation  ",
        "**AI Engine:** Hybrid AI Architecture — University of Moratuwa",
        "", sep, "",
        "## Zone 1 — PV Array Fault Diagnosis",
        "",
        f"| Field | Value |",
        f"|---|---|",
        f"| Fault Type | {r1['label'] if r1 else 'No data provided'} |",
    ]
    if r1 and r1.get("s2_label"):
        lines.append(f"| Stage 2 Localisation | {r1['s2_label']} |")
        lines.append(f"| Affected Array | {r1.get('stage2_array','—')} |")
        if r1.get("stage2_fault_type") == "Short-Circuit":
            lines.append(f"| R_LL | {r1.get('r_ll','—')} |")
            lines.append(f"| SC Severity | {r1.get('sc_severity','—')} |")
        else:
            lines.append(f"| R_LL | Not applicable (Open-Circuit) |")
    lines += ["", sep, "",
              "## Zone 2 — IGBT Inverter Fault Diagnosis", "",
              f"| Field | Value |", f"|---|---|",
              f"| Status | {r2['label'] if r2 else 'No data provided'} |",
              f"| Model Confidence | {_confidence(r2)} |",
              "", sep, "",
              "## Zone 3 — DC-Link Capacitor Health", ""]
    if r3:
        lines += [
            f"| Field | Value |", f"|---|---|",
            f"| Predicted ESR | {r3['esr']:.4f} Ω |",
            f"| Fault Threshold | 0.30 Ω |",
            f"| Degradation | {r3['deg']:.1f}% |",
            f"| Status | {r3['label']} |",
            f"| Irradiance | {r3['irr']:.0f} W/m² |",
            f"| T_ambient | {r3['tamb']:.1f} °C |",
            f"| V_avg | {r3['vavg']:.1f} V |",
            f"| V_ripple | {r3['vrip']:.4f} V |",
        ]
    else:
        lines.append("_No data provided_")
    if llm_text:
        lines += ["", sep, "",
                  "## Engineering Analysis", "",
                  _sanitize_report_text(llm_text)]
    lines += ["", sep,
              "",
              "_This report is an AI-assisted diagnostic aid. "
              "Final engineering decisions must be verified using site measurements and standard safety procedures._"]
    return "\n".join(lines)


# ════════════════════════════════════════════════════════════════
# EXPORT — DOCX
# ════════════════════════════════════════════════════════════════

def make_docx_report_file(r1, r2, r3, llm_text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin   = Inches(0.85)
        section.right_margin  = Inches(0.85)

    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PV Plant Fault Diagnosis Report")
    run.bold = True; run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x06, 0x3b, 0x2f)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Hybrid AI Architecture for Solar PV Fault Diagnosis  |  Generated: {ts}"
                ).font.size = Pt(9)

    doc.add_paragraph()

    # Helper: section heading
    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x0f, 0x76, 0x6e)

    # Helper: add table
    def add_styled_table(headers, rows):
        tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = "Table Grid"
        # Header row
        hdr = tbl.rows[0]
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = h
            set_cell_bg(cell, "0f766e")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                run.font.size = Pt(9)
        # Data rows
        for ri, row in enumerate(rows):
            tr = tbl.rows[ri + 1]
            fill = "f8fafc" if ri % 2 == 0 else "ffffff"
            for ci, val in enumerate(row):
                cell = tr.cells[ci]
                cell.text = str(val)
                set_cell_bg(cell, fill)
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)
        return tbl

    # Section 1 — Executive Summary
    add_heading("1. Executive Summary")
    z1_s = r1["label"] if r1 else "No data"
    z2_s = r2["label"] if r2 else "No data"
    z3_s = r3["label"] if r3 else "No data"
    overall = "FAULT DETECTED" if (
        (r1 and r1.get("cls") != 0) or
        (r2 and r2.get("cls") != 0) or
        (r3 and r3.get("fault"))
    ) else "NORMAL / MONITORING"
    add_styled_table(
        ["Item", "Result", "Confidence / Detail"],
        [
            ["System Status", overall, "Hybrid diagnostic output"],
            ["Zone 1 — PV Array", z1_s, _confidence(r1)],
            ["Zone 2 — IGBT Inverter", z2_s, _confidence(r2)],
            ["Zone 3 — DC-Link Capacitor", z3_s, "XGBoost ESR regression"],
        ]
    )

    # Section 2 — Zone detail
    doc.add_paragraph()
    add_heading("2. Zone-Specific Diagnostic Outputs")
    z1_detail = "No Stage 2 output"
    if r1 and r1.get("s2_label"):
        if r1.get("stage2_fault_type") == "Short-Circuit":
            z1_detail = (f"Array: {r1.get('stage2_array')} | R_LL: {r1.get('r_ll')} | "
                         f"Severity: {r1.get('sc_severity')}")
        else:
            z1_detail = f"Array: {r1.get('stage2_array')} | R_LL: Not applicable (OC fault)"
    z3_detail = (f"ESR: {r3['esr']:.4f} Ω | Deg: {r3['deg']:.1f}% | "
                 f"V_ripple: {r3['vrip']:.4f} V | T: {r3['tamb']:.1f} °C") if r3 else "No data"
    add_styled_table(
        ["Zone", "Primary Output", "Engineering Detail"],
        [
            ["Zone 1 — PV Array", z1_s, z1_detail],
            ["Zone 2 — IGBT Inverter", z2_s, "Alpha-beta current trajectory analysis"],
            ["Zone 3 — DC-Link Capacitor", z3_s, z3_detail],
        ]
    )

    # Section 3 — Engineering Analysis
    doc.add_paragraph()
    add_heading("3. Engineering Analysis")
    report_text = _sanitize_report_text(llm_text or _local_report(r1, r2, r3))
    for line in report_text.split("\n"):
        s = line.strip()
        if not s:
            doc.add_paragraph()
            continue
        m = re.match(r"^\*\*(\d+\.\s*.+?)\*\*$", s)
        if m:
            p = doc.add_paragraph(m.group(1))
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0x0f, 0x76, 0x6e)
                run.font.size = Pt(10)
            continue
        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            content = re.sub(r"\*\*(.*?)\*\*", r"\1", s[2:])
            p.add_run(content).font.size = Pt(9)
            continue
        p = doc.add_paragraph(re.sub(r"\*\*(.*?)\*\*", r"\1", s))
        p.runs[0].font.size = Pt(9) if p.runs else None

    # Section 4 — Verification note
    doc.add_paragraph()
    add_heading("4. Verification Note")
    note = doc.add_paragraph(
        "This report is an AI-assisted diagnostic aid. Final engineering decisions should be verified "
        "using electrical measurements, thermal inspection, site safety procedures, and the plant "
        "maintenance policy."
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════════
# EXPORT — PDF
# ════════════════════════════════════════════════════════════════

def _sanitize_report_text(text: str) -> str:
    text = "" if text is None else str(text)
    text = re.sub(
        r"\[(?:URGENT|CRITICAL|WITHIN\s*48H|SCHEDULED|IMMEDIATE|PRIORITY\s*\d+)\]\s*",
        "", text, flags=re.IGNORECASE
    )
    return text.strip()


def _pdf_safe(text) -> str:
    text = "" if text is None else str(text)
    for a, b in {
        "Ω": " Ohm", "≥": ">=", "≤": "<=", "–": "-", "—": "-",
        "α": "alpha", "β": "beta", "°": " deg", "•": "-",
        "→": "->", "×": "x", "²": "2",
        "⚠️": "", "🔴": "", "✅": "", "⚪": "", "📍": "",
        "\u2019": "'", "\u2018": "'", "\u201c": '"', "\u201d": '"',
    }.items():
        text = text.replace(a, b)
    # Strip any remaining non-ASCII to avoid ReportLab encoding errors
    return text.encode("ascii", errors="ignore").decode("ascii").strip()


def _safe_para(text, style):
    """Create a ReportLab Paragraph safely — falls back to plain text on error."""
    from reportlab.platypus import Paragraph
    import html
    safe = html.escape(_pdf_safe(text))
    try:
        return Paragraph(safe, style)
    except Exception:
        return Paragraph(_pdf_safe(text)[:200], style)


def _md_para(text, style):
    """Paragraph with basic **bold** and bullet conversion."""
    from reportlab.platypus import Paragraph
    import html
    text = _pdf_safe(text).strip()
    text = re.sub(r"^[-*]\s+", "• ", text)
    escaped = html.escape(text)
    escaped = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", escaped)
    try:
        return Paragraph(escaped, style)
    except Exception:
        return Paragraph(html.escape(_pdf_safe(text))[:200], style)


def make_pdf_report(r1, r2, r3, llm_text: str) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer,
        Table, TableStyle, HRFlowable
    )

    # ── Colour palette ───────────────────────────────────────────
    C_TEAL   = colors.HexColor("#0f766e")
    C_NAVY   = colors.HexColor("#083344")
    C_GRID   = colors.HexColor("#cbd5e1")
    C_ALT    = colors.HexColor("#f8fafc")
    C_DARK   = colors.HexColor("#0f172a")
    C_MUTED  = colors.HexColor("#334155")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=13*mm, leftMargin=13*mm,
        topMargin=13*mm,  bottomMargin=13*mm,
    )

    # ── Styles ───────────────────────────────────────────────────
    base = getSampleStyleSheet()
    base.add(ParagraphStyle("RTitle", parent=base["Title"],
        fontName="Helvetica-Bold", fontSize=17, leading=21,
        alignment=TA_CENTER, textColor=colors.HexColor("#063b2f"), spaceAfter=3))
    base.add(ParagraphStyle("RSub", parent=base["Normal"],
        fontSize=8, leading=10, alignment=TA_CENTER,
        textColor=colors.HexColor("#475569"), spaceAfter=6))
    base.add(ParagraphStyle("RSection", parent=base["Normal"],
        fontName="Helvetica-Bold", fontSize=11, leading=14,
        textColor=C_TEAL, spaceBefore=10, spaceAfter=5))
    base.add(ParagraphStyle("RBody", parent=base["Normal"],
        fontSize=8.5, leading=11.5, textColor=C_DARK,
        spaceAfter=4, wordWrap="CJK"))
    base.add(ParagraphStyle("RSmall", parent=base["Normal"],
        fontSize=7.5, leading=10, textColor=C_MUTED,
        spaceAfter=3, wordWrap="CJK"))

    # ── Shared table style helper ────────────────────────────────
    def base_table_style(header_color):
        return TableStyle([
            ("BACKGROUND",    (0, 0), (-1, 0), header_color),
            ("TEXTCOLOR",     (0, 0), (-1, 0), colors.white),
            ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0), (-1, 0), 8),
            ("FONTNAME",      (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE",      (0, 1), (-1, -1), 8),
            ("FONTNAME",      (0, 1), (0, -1), "Helvetica-Bold"),  # bold first col
            ("TEXTCOLOR",     (0, 1), (-1, -1), C_DARK),
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
            ("GRID",          (0, 0), (-1, -1), 0.3, C_GRID),
            ("LEFTPADDING",   (0, 0), (-1, -1), 5),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 5),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.white, C_ALT]),
        ])

    # ── Build content ────────────────────────────────────────────
    story = []
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    story.append(Paragraph("PV Plant Fault Diagnosis Report", base["RTitle"]))
    story.append(Paragraph(
        f"Hybrid AI Architecture for Solar PV Fault Diagnosis  |  Generated: {ts}",
        base["RSub"]
    ))
    story.append(HRFlowable(width="100%", thickness=1.0, color=C_TEAL))
    story.append(Spacer(1, 6))

    # ── Resolve result labels ────────────────────────────────────
    z1_s = _pdf_safe(r1["label"]) if r1 else "No data provided"
    z2_s = _pdf_safe(r2["label"]) if r2 else "No data provided"
    z3_s = _pdf_safe(r3["label"]) if r3 else "No data provided"
    z1_conf = _confidence(r1)
    z2_conf = _confidence(r2)
    overall = "FAULT DETECTED" if (
        (r1 and r1.get("cls") != 0) or
        (r2 and r2.get("cls") != 0) or
        (r3 and r3.get("fault"))
    ) else "NORMAL / MONITORING"

    # ── Section 1: Executive Summary ─────────────────────────────
    story.append(Paragraph("1. Executive Summary", base["RSection"]))
    t1_data = [
        ["Item",                      "Result",   "Model Evidence"],
        ["System Status",             overall,    "Hybrid diagnostic output"],
        ["Zone 1 — PV Array",         z1_s,       f"Confidence: {z1_conf}"],
        ["Zone 2 — IGBT Inverter",    z2_s,       f"Confidence: {z2_conf}"],
        ["Zone 3 — DC-Link Capacitor",z3_s,       "XGBoost ESR regression"],
    ]
    t1 = Table(t1_data, colWidths=[44*mm, 72*mm, 60*mm],
               repeatRows=1, splitByRow=1)
    t1.setStyle(base_table_style(C_TEAL))
    story.append(t1)
    story.append(Spacer(1, 6))

    # ── Section 2: Zone-Specific Outputs ─────────────────────────
    story.append(Paragraph("2. Zone-Specific Outputs", base["RSection"]))

    z1_detail = "Stage 2 localisation not run"
    if r1 and r1.get("s2_label"):
        if r1.get("stage2_fault_type") == "Short-Circuit":
            z1_detail = (
                f"Faulty array: {r1.get('stage2_array', 'Unknown')}  |  "
                f"R_LL = {r1.get('r_ll', 'N/A')}  |  "
                f"Severity: {r1.get('sc_severity', 'N/A')}"
            )
        else:
            z1_detail = (
                f"Faulty array: {r1.get('stage2_array', 'Unknown')}  |  "
                f"R_LL: Not applicable (open-circuit fault)"
            )

    z3_detail = "No data provided"
    if r3:
        z3_detail = (
            f"Predicted ESR = {r3['esr']:.4f} Ohm  |  "
            f"Degradation = {r3['deg']:.1f}%  |  "
            f"V_ripple = {r3['vrip']:.4f} V  |  "
            f"T_ambient = {r3['tamb']:.1f} deg C"
        )

    t2_data = [
        ["Zone",    "Primary Output", "Engineering Detail"],
        ["Zone 1",  z1_s,             _pdf_safe(z1_detail)],
        ["Zone 2",  z2_s,             "Alpha-beta current trajectory (Clarke transform)"],
        ["Zone 3",  z3_s,             _pdf_safe(z3_detail)],
    ]
    t2 = Table(t2_data, colWidths=[20*mm, 55*mm, 101*mm],
               repeatRows=1, splitByRow=1)
    t2.setStyle(base_table_style(C_NAVY))
    story.append(t2)
    story.append(Spacer(1, 6))

    # ── Section 3: Engineering Analysis ──────────────────────────
    story.append(Paragraph("3. Engineering Analysis", base["RSection"]))
    raw_analysis = _sanitize_report_text(llm_text or _local_report(r1, r2, r3))
    for line in raw_analysis.split("\n"):
        block = line.strip()
        if not block:
            story.append(Spacer(1, 3))
            continue
        # Skip GPT document title lines
        if re.match(r"^#+\s*Engineering", block, re.IGNORECASE):
            continue
        # Convert stray ## headings
        hm = re.match(r"^#{1,2}\s+(.+)", block)
        if hm:
            story.append(Paragraph(_pdf_safe(hm.group(1)), base["RSection"]))
            continue
        # Standalone bold title: **Title**
        bm = re.match(r"^\*\*([^*]+)\*\*$", block)
        if bm:
            story.append(Paragraph(_pdf_safe(bm.group(1)), base["RSection"]))
            continue
        story.append(_md_para(block, base["RBody"]))

    # ── Section 4: Verification Note ─────────────────────────────
    story.append(Spacer(1, 4))
    story.append(Paragraph("4. Verification Note", base["RSection"]))
    story.append(Paragraph(
        "This report is an AI-assisted diagnostic aid. "
        "Final engineering decisions must be verified using electrical "
        "measurements, thermal inspection, site safety procedures, "
        "and the plant maintenance policy.",
        base["RSmall"]
    ))

    doc.build(story)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════

def main():
    st.markdown('<div class="main-title">▣ PV-FAULT//ENGINEERING_DIAGNOSTIC_TERMINAL</div>', unsafe_allow_html=True)
    st.markdown('<div class="main-sub">HYBRID AI ARCHITECTURE · 5 kW GRID-CONNECTED PV SYSTEM · UOM ELECTRICAL ENGINEERING</div>', unsafe_allow_html=True)

    M = load_models()
    if not M.get("ok"):
        st.error(f"Model loading failed: {M.get('err')}")
        st.info("Ensure all files are in the `models/` folder. See README.md.")
        return

    defaults = {"r1":None,"r2":None,"r3":None,"df1":None,"df1a":None,
                "df2":None,"df3":None,"llm":""}
    for k, v in defaults.items():
        if k not in st.session_state: st.session_state[k] = v

    # ── Overall status strip ─────────────────────────────────────
    oc1, oc2, oc3 = st.columns(3)
    for col, r, zone in [(oc1,st.session_state.r1,1),(oc2,st.session_state.r2,2),(oc3,st.session_state.r3,3)]:
        c, lbl = zone_status(r, zone)
        zone_names = {1:"Zone 1 — PV Array",2:"Zone 2 — IGBT Inverter",3:"Zone 3 — Capacitor"}
        col.markdown(badge(f"{zone_names[zone]}: {lbl}",c), unsafe_allow_html=True)

    active_faults = sum([
        1 if st.session_state.r1 and st.session_state.r1["cls"]!=0 else 0,
        1 if st.session_state.r2 and st.session_state.r2["cls"]!=0 else 0,
        1 if st.session_state.r3 and st.session_state.r3["fault"] else 0,
    ])
    tc = "#ef4444" if active_faults else "#22c55e"
    st.markdown(terminal_panel("system bus",{
        "MODE":"LIVE INFERENCE / DEMO",
        "MODEL STACK":"Zone1 LSTM · Zone2 CNN-1D · Zone3 XGBoost",
        "ACTIVE FAULTS":active_faults,
        "TIMESTAMP":datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    },tc), unsafe_allow_html=True)

    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    z1, z2, z3 = st.columns(3, gap="medium")

    # ══════════════════════════════════════════════════
    # ZONE 1
    # ══════════════════════════════════════════════════
    with z1:
        st.markdown('<div class="zone-title">🌞 Zone 1 — PV Array (LSTM)</div>', unsafe_allow_html=True)
        a, b, c, d = st.columns(4)
        if a.button("Healthy", key="z1h"):  st.session_state.df1 = sample("z1_healthy.csv")
        if b.button("SC Fault",key="z1sc"): st.session_state.df1 = sample("z1_sc.csv")
        if c.button("OC Fault",key="z1oc"): st.session_state.df1 = sample("z1_oc.csv")
        if d.button("DG Fault",key="z1dg"): st.session_state.df1 = sample("z1_dg.csv")
        uf1 = st.file_uploader("Upload IV Sweep CSV",type="csv",key="uf1",help="Columns: V, I, P, Irr, T")
        if uf1:
            df_temp = clean_df(pd.read_csv(uf1, encoding="utf-8-sig"))
            missing = [c for c in ["V","I","P","Irr","T"] if c not in df_temp.columns]
            if missing:
                st.error(f"Missing columns: {', '.join(missing)}")
                if "V_1" in df_temp.columns: st.info("💡 Looks like per-array CSV — upload in Stage 2 section.")
                st.session_state.df1 = None
            else: st.session_state.df1 = df_temp
        if st.session_state.df1 is not None:
            st.plotly_chart(plot_iv(st.session_state.df1), use_container_width=True, config=_PLOT_CFG)
            if st.button("🔍 Run Zone 1 Diagnosis",key="r1btn",type="primary",use_container_width=True):
                with st.spinner("Analysing IV curve…"):
                    try: st.session_state.r1 = run_z1(st.session_state.df1,st.session_state.df1a,M); st.rerun()
                    except Exception as e: st.error(f"Zone 1 error: {e}")
        if st.session_state.r1:
            r = st.session_state.r1
            c_col, _ = zone_status(r, 1)
            st.markdown(badge(r["label"],c_col), unsafe_allow_html=True)
            st.markdown(prob_bars(Z1_S1_LABELS,r["probs"],r["cls"]), unsafe_allow_html=True)
            if r["cls"] in [2, 3]:
                st.markdown("**Stage 2 — Array Localisation**")
                st.markdown('<div class="demo-caption">Select per-array sample or upload CSV</div>', unsafe_allow_html=True)
                if r["cls"] == 2:
                    sc_cols = st.columns(4)
                    sc_opts = [("A1 0.05Ω","z1_sc_array1_rll_0p05.csv"),("A1 10Ω","z1_sc_array1_rll_10p0.csv"),
                               ("A1 20Ω","z1_sc_array1_rll_20p0.csv"),("A1 30Ω","z1_sc_array1_rll_30p0.csv"),
                               ("A2 0.05Ω","z1_sc_array2_rll_0p05.csv"),("A2 10Ω","z1_sc_array2_rll_10p0.csv"),
                               ("A2 20Ω","z1_sc_array2_rll_20p0.csv"),("A2 30Ω","z1_sc_array2_rll_30p0.csv")]
                    for i,(txt,fname) in enumerate(sc_opts):
                        if sc_cols[i%4].button(txt,key=f"z1s2_{fname}",use_container_width=True):
                            st.session_state.df1a = sample(fname); st.session_state.llm = ""
                elif r["cls"] == 3:
                    oc_cols = st.columns(2)
                    if oc_cols[0].button("OC Array 1",key="z1oc_a1",use_container_width=True):
                        st.session_state.df1a = sample("z1_oc_array1.csv"); st.session_state.llm = ""
                    if oc_cols[1].button("OC Array 2",key="z1oc_a2",use_container_width=True):
                        st.session_state.df1a = sample("z1_oc_array2.csv"); st.session_state.llm = ""
                uf1a = st.file_uploader("Upload Per-Array CSV",type="csv",key="uf1a",
                                        help="Columns: V_1, I_1, P_1, V_2, I_2, P_2, Irr, T")
                if uf1a:
                    df_temp = clean_df(pd.read_csv(uf1a, encoding="utf-8-sig"))
                    missing = [c for c in ["V_1","I_1","P_1","V_2","I_2","P_2","Irr","T"] if c not in df_temp.columns]
                    if missing: st.error(f"Missing Stage 2 columns: {', '.join(missing)}"); st.session_state.df1a = None
                    else: st.session_state.df1a = df_temp
                if st.session_state.df1a is not None:
                    if st.button("🔍 Localise Fault",key="r1s2btn",type="secondary",use_container_width=True):
                        with st.spinner("Localising…"):
                            try: st.session_state.r1 = run_z1(st.session_state.df1,st.session_state.df1a,M); st.rerun()
                            except Exception as e: st.error(f"Stage 2 error: {e}")
            if r.get("s2_label"):
                st.markdown(badge(f"📍 {r['s2_label']}","#8b5cf6"), unsafe_allow_html=True)
                detail = {"FAULT ROUTE":r.get("stage2_fault_type","Unknown"),"LOCALISATION":r.get("stage2_array","Unknown")}
                if r.get("stage2_fault_type") == "Short-Circuit":
                    detail["R_LL OUTPUT"] = r.get("r_ll","—")
                    detail["SC SEVERITY"] = r.get("sc_severity","—")
                else: detail["R_LL OUTPUT"] = "N/A for open-circuit"
                st.markdown(terminal_panel("zone 1 stage-2 decoded output",detail,"#8b5cf6"), unsafe_allow_html=True)

    # ══════════════════════════════════════════════════
    # ZONE 2
    # ══════════════════════════════════════════════════
    with z2:
        st.markdown('<div class="zone-title">⚡ Zone 2 — IGBT Inverter (CNN-1D)</div>', unsafe_allow_html=True)
        st.markdown('<div class="demo-caption">Demo samples — all 7 inverter classes</div>', unsafe_allow_html=True)
        z2r1 = st.columns(4); z2r2 = st.columns(3)
        z2_samples = [("Healthy","z2_healthy.csv"),("T1 Fault","z2_t1_fault.csv"),
                      ("T2 Fault","z2_t2_fault.csv"),("T3 Fault","z2_t3_fault.csv"),
                      ("T4 Fault","z2_t4_fault.csv"),("T5 Fault","z2_t5_fault.csv"),
                      ("T6 Fault","z2_t6_fault.csv")]
        for i,(txt,fname) in enumerate(z2_samples):
            col = z2r1[i] if i < 4 else z2r2[i-4]
            if col.button(txt,key=f"z2_{fname}",use_container_width=True):
                st.session_state.df2 = sample(fname); st.session_state.r2 = None; st.session_state.llm = ""
        uf2 = st.file_uploader("Upload Current Trajectory CSV",type="csv",key="uf2",
                                help="Columns: I_alpha, I_beta, Irradiance, Temperature (≥32 rows)")
        if uf2:
            df_temp = clean_df(pd.read_csv(uf2, encoding="utf-8-sig"))
            missing = [c for c in ["I_alpha","I_beta","Irradiance","Temperature"] if c not in df_temp.columns]
            if missing: st.error(f"Missing columns: {', '.join(missing)}"); st.session_state.df2 = None
            else: st.session_state.df2 = df_temp
        if st.session_state.df2 is not None:
            if len(st.session_state.df2) < WIN_Z2:
                st.warning(f"CSV needs ≥{WIN_Z2} rows. Found {len(st.session_state.df2)}.")
            else:
                st.plotly_chart(plot_ab(st.session_state.df2), use_container_width=True, config=_PLOT_CFG)
                if st.button("🔍 Run Zone 2 Diagnosis",key="r2btn",type="primary",use_container_width=True):
                    with st.spinner("Analysing αβ trajectory…"):
                        try: st.session_state.r2 = run_z2(st.session_state.df2,M); st.rerun()
                        except Exception as e: st.error(f"Zone 2 error: {e}")
        if st.session_state.r2:
            r = st.session_state.r2
            c_col,_ = zone_status(r,2)
            st.markdown(badge(r["label"],c_col), unsafe_allow_html=True)
            st.markdown(prob_bars(Z2_LABELS,r["probs"],r["cls"]), unsafe_allow_html=True)

    # ══════════════════════════════════════════════════
    # ZONE 3
    # ══════════════════════════════════════════════════
    with z3:
        st.markdown('<div class="zone-title">🔋 Zone 3 — DC-Link Capacitor (XGBoost)</div>', unsafe_allow_html=True)
        a, b = st.columns(2)
        if a.button("Healthy",  key="z3h"): st.session_state.df3 = sample("z3_healthy.csv")
        if b.button("Degraded", key="z3d"): st.session_state.df3 = sample("z3_degraded.csv")
        uf3 = st.file_uploader("Upload Measurement CSV",type="csv",key="uf3",
                                help="Columns: Irradiance, T_ambient, I_rms_high, V_avg, V_ripple")
        if uf3:
            df_temp = clean_df(pd.read_csv(uf3, encoding="utf-8-sig"))
            missing = [c for c in ["Irradiance","T_ambient","I_rms_high","V_avg","V_ripple"] if c not in df_temp.columns]
            if missing: st.error(f"Missing columns: {', '.join(missing)}"); st.session_state.df3 = None
            else: st.session_state.df3 = df_temp
        if st.session_state.df3 is not None:
            row = st.session_state.df3.iloc[0]
            ma, mb = st.columns(2)
            ma.metric("Irradiance", f"{row.get('Irradiance',0):.0f} W/m²")
            ma.metric("V_avg",      f"{row.get('V_avg',0):.1f} V")
            mb.metric("T_ambient",  f"{row.get('T_ambient',0):.1f} °C")
            mb.metric("V_ripple",   f"{row.get('V_ripple',0):.4f} V")
            if st.button("🔍 Run Zone 3 Diagnosis",key="r3btn",type="primary",use_container_width=True):
                with st.spinner("Estimating ESR…"):
                    try: st.session_state.r3 = run_z3(st.session_state.df3,M); st.rerun()
                    except Exception as e: st.error(f"Zone 3 error: {e}")
        if st.session_state.r3:
            r = st.session_state.r3
            st.plotly_chart(plot_gauge(r["esr"],r["deg"]), use_container_width=True, config=_PLOT_CFG)
            c_col,_ = zone_status(r,3)
            st.markdown(badge(r["label"],c_col), unsafe_allow_html=True)
            st.markdown(terminal_panel("zone 3 health-index output",{
                "PREDICTED ESR":f'{r["esr"]:.4f} Ω',"AGEING BAND":"0.15–0.30 Ω",
                "FAULT THRESHOLD":"≥ 0.30 Ω","DEGRADATION %":f'{r["deg"]:.1f}%',
                "INTERPRETATION":"Hard fault" if r["fault"] else ("Ageing/warning" if r["esr"]>0.15 else "Healthy"),
            },c_col), unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════════
    # ENGINEERING ANALYSIS — Modern clean layout
    # ════════════════════════════════════════════════════════════
    any_result = any([st.session_state.r1, st.session_state.r2, st.session_state.r3])
    if any_result:
        st.markdown("---")

        # ── Section header ───────────────────────────────────────
        st.markdown("""
        <div class="analysis-wrapper">
            <div class="analysis-top-row">
                <div>
                    <div class="analysis-title">Engineering Analysis Report</div>
                    <div class="analysis-subtitle">AI-assisted fault assessment · GPT-4o-mini · University of Moratuwa</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        gen_col, spacer, export_col = st.columns([2, 0.1, 1])

        with gen_col:
            # ── Generate button ───────────────────────────────────
            if st.button("⚙ Generate Engineering Analysis",
                         type="primary", use_container_width=True):
                with st.spinner("Generating analysis…"):
                    st.session_state.llm = gpt_explain(
                        st.session_state.r1,
                        st.session_state.r2,
                        st.session_state.r3,
                    )

            # ── Report display ────────────────────────────────────
            if st.session_state.llm:
                rendered = render_report_html(st.session_state.llm)
                st.markdown(
                    f'<div class="report-card-outer">{rendered}</div>',
                    unsafe_allow_html=True,
                )

        with export_col:
            st.markdown('<div class="export-panel">', unsafe_allow_html=True)
            st.markdown('<div class="export-panel-title">Export Report</div>', unsafe_allow_html=True)

            fmt = st.selectbox(
                "Format",
                ["PDF (.pdf)", "Markdown (.md)", "DOCX (.docx)"],
                label_visibility="collapsed",
            )

            ts_str = datetime.now().strftime("%Y%m%d_%H%M")

            if fmt == "PDF (.pdf)":
                with st.spinner("Building PDF…"):
                    try:
                        pdf = make_pdf_report(
                            st.session_state.r1, st.session_state.r2,
                            st.session_state.r3, st.session_state.llm,
                        )
                        st.download_button(
                            label="⬇ Download PDF",
                            data=pdf,
                            file_name=f"pv_report_{ts_str}.pdf",
                            mime="application/pdf",
                            use_container_width=True,
                        )
                    except Exception as e:
                        st.error(f"PDF error: {e}")

            elif fmt == "Markdown (.md)":
                md = make_markdown_report(
                    st.session_state.r1, st.session_state.r2,
                    st.session_state.r3, st.session_state.llm,
                )
                st.download_button(
                    label="⬇ Download Markdown",
                    data=md.encode("utf-8"),
                    file_name=f"pv_report_{ts_str}.md",
                    mime="text/markdown",
                    use_container_width=True,
                )

            elif fmt == "DOCX (.docx)":
                if st.button("⬇ Download DOCX", use_container_width=True, type="primary"):
                    with st.spinner("Building DOCX…"):
                        try:
                            docx_bytes = make_docx_report_file(
                                st.session_state.r1, st.session_state.r2,
                                st.session_state.r3, st.session_state.llm,
                            )
                            st.download_button(
                                label="📝 Click to save DOCX",
                                data=docx_bytes,
                                file_name=f"pv_report_{ts_str}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                            )
                        except Exception as e:
                            st.error(f"DOCX error: {e}. Run: pip install python-docx")

            st.markdown("&nbsp;")
            if st.button("🔄 Reset All Zones", use_container_width=True):
                for k in defaults: st.session_state[k] = defaults[k]
                st.rerun()

            st.markdown('</div>', unsafe_allow_html=True)


if __name__ == "__main__":
    main()